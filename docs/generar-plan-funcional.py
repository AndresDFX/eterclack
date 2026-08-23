"""
Genera el plan de pruebas funcionales en .docx.

El documento va dirigido a quien prueba a mano, no a quien programa. Cada caso
trae pasos, resultado esperado y un bloque de registro para anotar qué pasó de
verdad: sin ese espacio, las observaciones se pierden en un chat y no llegan a
convertirse en correcciones.

    python docs/generar-plan-funcional.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Marca ────────────────────────────────────────────────────────
INK = RGBColor(0x1D, 0x1D, 0x1B)
LIME_HEX = 'E7E226'
BAND_HEX = 'F2F2EE'   # banda del bloque de registro
FIELD_HEX = 'FFFFFF'  # celdas para escribir
GREY = RGBColor(0x6A, 0x6A, 0x66)
DANGER = RGBColor(0xB0, 0x3A, 0x2E)

CAJA = '☐'

doc = Document()

for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(1.8)
    s.left_margin = s.right_margin = Cm(1.8)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(6)


# ── Utilidades ───────────────────────────────────────────────────
def sombrear(celda, hex_color):
    tc = celda._tc.get_or_add_tcPr()
    sombra = OxmlElement('w:shd')
    sombra.set(qn('w:val'), 'clear')
    sombra.set(qn('w:fill'), hex_color)
    tc.append(sombra)


def texto_celda(celda, texto, size=9, bold=False, color=None, italic=False):
    celda.text = ''
    p = celda.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def no_partir(tabla_obj):
    """Mantiene el bloque de un caso entero en la misma página.

    Un caso partido entre dos hojas obliga a pasar página para anotar lo que
    acabas de ver: la observación se pierde. Se marca cada fila como
    indivisible y se encadena con la siguiente, salvo la última.
    """
    filas = tabla_obj.rows
    for indice, fila in enumerate(filas):
        tr = fila._tr.get_or_add_trPr()
        # El esquema de OOXML exige que `cantSplit` vaya ANTES de `trHeight`.
        # Puesto después, Word lo ignora en silencio y la fila se parte igual.
        no_split = OxmlElement('w:cantSplit')
        no_split.set(qn('w:val'), 'true')
        tr.insert(0, no_split)
        if indice < len(filas) - 1:
            for celda in fila.cells:
                for par in celda.paragraphs:
                    par.paragraph_format.keep_with_next = True


def titulo(texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for r in h.runs:
        r.font.color.rgb = INK
        r.font.name = 'Calibri'
    return h


def parrafo(texto, gris=False, negrita=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(size)
    r.bold = negrita
    if gris:
        r.font.color.rgb = GREY
    return p


def nota(texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(texto)
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = GREY
    return p


def tabla(cabeceras, filas, anchos=None, size=9):
    t = doc.add_table(rows=1, cols=len(cabeceras))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(cabeceras):
        texto_celda(t.rows[0].cells[i], h, size=size, bold=True)
        sombrear(t.rows[0].cells[i], LIME_HEX)
    for fila in filas:
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            texto_celda(celdas[i], str(v), size=size)
    if anchos:
        for fila in t.rows:
            for i, w in enumerate(anchos):
                fila.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# ════════════════════════════════════════════════════════════════
#  Portada
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run('ETERCLACK')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = INK

p = doc.add_paragraph()
r = p.add_run('FOTOGRAFÍA Y VIDEO')
r.font.size = Pt(8)
r.font.color.rgb = GREY

doc.add_paragraph()
titulo('Plan de pruebas funcionales', 0)
parrafo(
    'Guion para probar la plataforma a mano. Cada caso indica quién lo ejecuta, qué hace, qué '
    'debe pasar, y deja espacio para anotar qué pasó de verdad. Los casos marcados como '
    'NEGATIVOS deben FALLAR: si pasan, hay un defecto.',
    gris=True,
)

parrafo('Datos de la sesión de pruebas', negrita=True)
parrafo('Llena esto antes de empezar. Sin el contexto, una observación no se puede reproducir.',
        gris=True, size=9)

t = doc.add_table(rows=0, cols=4)
t.style = 'Table Grid'
for etiqueta_a, etiqueta_b in [
    ('Quién prueba', 'Fecha'),
    ('Entorno (local / Render)', 'URL usada'),
    ('Navegador y versión', 'Dispositivo'),
    ('Versión probada (commit)', 'Duración'),
]:
    fila = t.add_row()
    fila.height = Cm(0.85)
    fila.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    texto_celda(fila.cells[0], etiqueta_a, bold=True)
    sombrear(fila.cells[0], BAND_HEX)
    texto_celda(fila.cells[1], '')
    texto_celda(fila.cells[2], etiqueta_b, bold=True)
    sombrear(fila.cells[2], BAND_HEX)
    texto_celda(fila.cells[3], '')
    fila.cells[0].width = Cm(4.4)
    fila.cells[1].width = Cm(4.6)
    fila.cells[2].width = Cm(4)
    fila.cells[3].width = Cm(4.4)
doc.add_paragraph()

tabla(
    ['Campo', 'Valor'],
    [
        ['Alcance', 'Fases 0 a 4: identidad, descubrimiento, calendario, reserva, contrato, PWA'],
        ['Fuera de alcance', 'Galerías y entrega, pagos con Wompi, dispersión, correo propio'],
        ['Casos totales', '79, de los cuales 24 son negativos'],
        ['Duración estimada', '2 horas el recorrido completo'],
    ],
    anchos=[4, 13],
)

# ════════════════════════════════════════════════════════════════
titulo('1. Cómo llenar este documento')

parrafo(
    'Bajo cada caso hay un bloque de registro. Márcalo siempre, incluso cuando todo va bien: '
    'un caso sin marcar no se distingue de uno sin probar.'
)

t = doc.add_table(rows=0, cols=2)
t.style = 'Table Grid'
for campo, explicacion in [
    ('Resultado',
     'OK si hizo exactamente lo esperado. Falla si no. Bloqueado si no pudiste llegar a probarlo '
     'porque algo anterior no funciona. Un caso a medias es Falla, no OK.'),
    ('Severidad',
     'Solo si falla. Bloqueante: impide seguir o pierde datos. Alta: rompe una función. '
     'Media: molesta pero hay forma de seguir. Baja: cosmético.'),
    ('Observaciones',
     'Qué viste exactamente, no una interpretación. «El botón queda deshabilitado» sirve; '
     '«no funciona» no. Si es visual, di en qué pantalla y ancho.'),
    ('Qué corregir',
     'Tu propuesta, si la tienes. Puede quedar vacío: es información para quien arregla, '
     'no una obligación de quien prueba.'),
]:
    fila = t.add_row()
    texto_celda(fila.cells[0], campo, bold=True)
    sombrear(fila.cells[0], BAND_HEX)
    texto_celda(fila.cells[1], explicacion)
    fila.cells[0].width = Cm(3.4)
    fila.cells[1].width = Cm(14)
doc.add_paragraph()

nota(
    'Al final del documento hay un registro consolidado de defectos. Copia allí solo los casos '
    'que fallaron: es la lista que se convierte en trabajo de corrección.'
)

# ════════════════════════════════════════════════════════════════
titulo('2. Antes de empezar')

parrafo('Direcciones', negrita=True)
tabla(
    ['Qué', 'Local', 'Render'],
    [
        ['Aplicación', 'http://localhost:5173', 'https://<tu-servicio>.onrender.com'],
        ['Correos recibidos', 'http://localhost:8025 (Mailpit)', 'No hay: el plan gratuito bloquea SMTP'],
        ['Salud del servicio', 'http://localhost:3000/health', '<tu-servicio>/health'],
        ['Salud de la base', 'http://localhost:3000/health/db', '<tu-servicio>/health/db'],
    ],
    anchos=[4, 6.5, 6.5],
)

nota(
    'En Render el servicio gratuito se duerme tras 15 minutos sin uso: la primera petición puede '
    'tardar entre 30 y 60 segundos. No es un fallo, y no debe registrarse como tal.'
)
nota(
    'En Render tampoco hay correo: el plan gratuito bloquea los puertos SMTP. Los casos CP-04, '
    'CP-05, CP-08 y CP-09 solo se pueden probar en local. Márcalos como Bloqueado.'
)

parrafo('Credenciales', negrita=True)
parrafo('Todas usan la misma contraseña: ', gris=True)
p = doc.add_paragraph()
r = p.add_run('Eterclack123*')
r.bold = True
r.font.size = Pt(12)
r.font.name = 'Consolas'

tabla(
    ['Correo', 'Rol', 'Estado', 'Para qué sirve'],
    [
        ['admin@eterclack.test', 'Administración', 'Activo', 'Aprobar, rechazar y suspender fotógrafos'],
        ['maria@eterclack.test', 'Fotógrafa', 'Aprobada · con agenda', 'Recibir reservas, publicar calendario'],
        ['andres@eterclack.test', 'Fotógrafo', 'Aprobado · con agenda', 'Segundo fotógrafo para comparar'],
        ['laura@eterclack.test', 'Fotógrafa', 'Aprobada · con agenda', 'Tercer fotógrafo'],
        ['carlos@eterclack.test', 'Fotógrafo', 'PENDIENTE de revisión', 'Probar el flujo de aprobación'],
        ['sofia@eterclack.test', 'Fotógrafa', 'RECHAZADA', 'Probar el motivo de rechazo y la repostulación'],
        ['juliana@eterclack.test', 'Cliente', 'Verificada', 'Reservar citas'],
        ['daniel@eterclack.test', 'Cliente', 'Verificada', 'Probar que no ve datos de Juliana'],
        ['valentina@eterclack.test', 'Cliente', 'Verificada', 'Tercer cliente'],
    ],
    anchos=[5, 3.2, 4, 5],
)

nota(
    'Los tres estados de fotógrafo (aprobado, pendiente, rechazado) están sembrados a propósito: '
    'permiten probar el ciclo de aprobación sin crear cuentas nuevas.'
)

# ════════════════════════════════════════════════════════════════
CASOS = [
    (
        '3. Cuenta y acceso',
        'Cliente nuevo',
        [
            ('CP-01', 'Registro de cliente',
             '1. Abre la aplicación y pulsa «Crear cuenta».\n'
             '2. Elige «Contratar».\n'
             '3. Llena nombre, un correo nuevo y una contraseña de 8+ caracteres.\n'
             '4. Marca la casilla de términos y envía.',
             'Entra a la aplicación. Aparece un aviso amarillo arriba pidiendo confirmar el correo.'),

            ('CP-02', 'Contraseña muy corta',
             'Repite CP-01 pero con una contraseña de 5 caracteres.',
             'NEGATIVO. No deja crear la cuenta. Bajo el campo aparece «La contraseña debe tener '
             'al menos 8 caracteres».'),

            ('CP-03', 'Correo repetido',
             'Intenta registrarte con juliana@eterclack.test.',
             'NEGATIVO. Mensaje indicando que ese correo ya tiene cuenta. No se crea un segundo usuario.'),

            ('CP-04', 'Verificación por correo',
             '1. Abre Mailpit (local) o tu bandeja.\n'
             '2. Busca «Confirma tu correo en EterClack».\n'
             '3. Pulsa «Confirmar mi correo».',
             'El aviso amarillo desaparece. El correo llega con los colores de la marca: fondo '
             'oscuro, botón amarillo, esquinas de visor.'),

            ('CP-05', 'Reusar el enlace de verificación',
             'Vuelve a abrir el mismo enlace del correo.',
             'NEGATIVO. «El enlace no es válido o ya venció». Un enlace se usa una sola vez.'),

            ('CP-06', 'Ingreso correcto',
             'Cierra sesión e ingresa con juliana@eterclack.test.',
             'Entra y aparece su nombre arriba a la derecha.'),

            ('CP-07', 'Contraseña incorrecta',
             'Ingresa con juliana@eterclack.test y la contraseña «loquesea».',
             'NEGATIVO. «Correo o contraseña incorrectos». Debe ser el MISMO mensaje que con un '
             'correo inexistente: la plataforma no revela qué correos están registrados.'),

            ('CP-08', 'Recuperar contraseña',
             '1. En el ingreso, pulsa «¿Olvidaste tu contraseña?».\n'
             '2. Escribe juliana@eterclack.test.\n'
             '3. Abre el correo y crea una contraseña nueva.\n'
             '4. Ingresa con la nueva.',
             'Funciona con la nueva contraseña. La anterior deja de servir.'),

            ('CP-09', 'Recuperar con correo inexistente',
             'Repite CP-08 con noexiste@ejemplo.com.',
             'Muestra el MISMO mensaje de confirmación que con un correo real. No llega ningún '
             'correo. Así no se puede averiguar quién tiene cuenta.'),
        ],
    ),
    (
        '4. Descubrimiento',
        'Cualquiera, sin necesidad de cuenta',
        [
            ('CP-10', 'Ver la portada',
             'Abre la aplicación sin iniciar sesión.',
             'Se ve la portada con «Tu momento, para siempre», las especialidades y los cuatro pasos.'),

            ('CP-11', 'Buscar fotógrafos',
             'Pulsa «Fotógrafos» en el menú.',
             'Aparecen 3 fotógrafos aprobados, cada uno con sus etiquetas y su precio desde. '
             'Con bucket configurado, además tres fotos de su portafolio.'),

            ('CP-12', 'Filtrar por especialidad',
             'En el filtro de especialidad elige «Bodas».',
             'Queda solo María Gómez. El contador dice «1 fotógrafo».'),

            ('CP-13', 'Filtrar por zona',
             'Limpia los filtros y elige la zona «Medellín».',
             'Queda solo Andrés Rueda.'),

            ('CP-14', 'Filtrar por presupuesto',
             'Limpia y pulsa «Hasta $500.000».',
             'Queda solo Andrés Rueda ($450.000). Los de mayor precio desaparecen.'),

            ('CP-15', 'Combinar filtros sin resultados',
             'Elige especialidad «Bodas» y zona «Medellín» a la vez.',
             'Aparece «Nada por aquí» con un botón para limpiar los filtros. No una pantalla en blanco.'),

            ('CP-16', 'Ver la ficha',
             'Limpia los filtros y abre María Gómez.',
             'Se ven su biografía, los TRES productos con precio y contenido, y el portafolio.'),

            ('CP-17', 'Fotógrafo pendiente no aparece',
             'Busca «Carlos» en el buscador.',
             'NEGATIVO. Cero resultados. Carlos está pendiente de aprobación y no debe ser visible.'),

            ('CP-18', 'URL directa a un pendiente',
             'Escribe en el navegador la dirección .../fotografos/carlos-duarte',
             'NEGATIVO. «Ese fotógrafo no está disponible». No se puede saltar la aprobación '
             'conociendo la dirección.'),
        ],
    ),
    (
        '5. Reserva de cita',
        'Cliente: juliana@eterclack.test',
        [
            ('CP-19', 'Reservar requiere cuenta',
             'Sin iniciar sesión, abre una ficha y pulsa «Solicitar fecha».',
             'Aparece una pantalla pidiendo ingresar o crear cuenta. No deja continuar.'),

            ('CP-20', 'Reservar requiere correo verificado',
             'Con la cuenta creada en CP-01 SIN verificar, intenta reservar.',
             'NEGATIVO. Pide confirmar el correo antes de apartar una fecha.'),

            ('CP-21', 'Los tres productos',
             'Ingresa como Juliana, abre María Gómez y pulsa «Solicitar fecha».',
             'Paso 01 muestra exactamente tres productos: Económico, Medio y Premium. El del '
             'medio lleva la etiqueta «Más elegido». Cada uno muestra qué incluye, horas, fotos '
             'y días de entrega.'),

            ('CP-22', 'Calendario con disponibilidad',
             'Mira el paso 02.',
             'Solo los días con cupo se pueden pulsar. Los demás están apagados. Los días con '
             'cupo llevan puntos amarillos abajo (uno por franja).'),

            ('CP-23', 'Elegir franja',
             'Pulsa un día con cupo.',
             'A la derecha aparecen las franjas de ese día: Mañana, Tarde o Jornada completa.'),

            ('CP-24', 'No se puede reservar sin elegir todo',
             'Elige solo el producto, sin fecha.',
             'El botón «Apartar esta fecha» queda deshabilitado. El resumen dice «Sin elegir».'),

            ('CP-25', 'Reserva completa',
             '1. Elige el producto Medio.\n'
             '2. Elige un día y una franja.\n'
             '3. Escribe una nota (opcional).\n'
             '4. Pulsa «Apartar esta fecha».',
             'Va a la pantalla de la orden. El código empieza por ETC-. El estado dice «Falta '
             'aceptar contrato». El precio es exactamente el del producto elegido.'),

            ('CP-26', 'La fecha desaparece del calendario',
             'Vuelve a la ficha de María y entra otra vez a reservar.',
             'La franja que acabas de tomar ya NO aparece disponible.'),

            ('CP-27', 'Dos personas, la misma fecha',
             '1. Abre otra ventana en modo incógnito.\n'
             '2. Ingresa como daniel@eterclack.test.\n'
             '3. Intenta reservar exactamente la misma franja.',
             'NEGATIVO. «Alguien acaba de tomar esa fecha. Elige otra del calendario.» '
             'Solo una reserva puede ganar.'),

            ('CP-28', 'Correos de la reserva',
             'Revisa la bandeja. Solo aplica en local.',
             'Dos correos: a la fotógrafa «Nueva cita: …» con fecha, franja, producto y valor; '
             'al cliente «Reservamos tu fecha» avisando que tiene 24 horas para confirmar.'),
        ],
    ),
    (
        '6. Contrato',
        'Cliente: juliana@eterclack.test',
        [
            ('CP-29', 'Leer el contrato',
             'En la pantalla de la orden, baja hasta «Contrato».',
             'El texto tiene los datos reales resueltos: nombres de las dos partes, fecha, lugar, '
             'paquete y valor. NO deben quedar cosas como {{cliente}} o {{valor}}.'),

            ('CP-30', 'No se acepta sin marcar',
             'Escribe tu nombre pero no marques la casilla.',
             'El botón «Aceptar contrato» sigue deshabilitado.'),

            ('CP-31', 'Aceptar el contrato',
             '1. Escribe tu nombre completo.\n'
             '2. Marca la casilla.\n'
             '3. Pulsa «Aceptar contrato».',
             'Aparece «Evidencia de aceptación» con el nombre que escribiste, la fecha y hora, y '
             'la versión de la plantilla. El estado de la orden pasa a «Pago pendiente».'),

            ('CP-32', 'No se acepta dos veces',
             'Recarga la página e intenta aceptar de nuevo.',
             'NEGATIVO. Ya no aparece el formulario, solo la evidencia. El contrato aceptado no se toca.'),

            ('CP-33', 'Correo de contrato',
             'Revisa la bandeja de ambas partes. Solo aplica en local.',
             'Las dos reciben «Contrato aceptado · Orden ETC-…».'),
        ],
    ),
    (
        '7. Panel del fotógrafo',
        'Fotógrafa: maria@eterclack.test',
        [
            ('CP-34', 'Ver la cita recibida',
             'Ingresa como María y entra a «Mis citas».',
             'Aparece la reserva de Juliana. Muestra «Recibes» con el valor NETO, ya descontada '
             'la comisión del 15 %.'),

            ('CP-35', 'La cuenta cuadra',
             'Abre el detalle de la orden y suma.',
             'Comisión + lo que recibe = valor total, exacto. Sin céntimos perdidos.'),

            ('CP-36', 'Ver el calendario',
             'Entra a «Calendario».',
             'Se ven los días publicados, los que tienen cita (en verde) y el conteo de cada uno. '
             'A la derecha aparece la próxima cita con el nombre del cliente.'),

            ('CP-37', 'Publicar disponibilidad',
             '1. Elige «Jornada completa» a la derecha.\n'
             '2. Pulsa 3 días futuros en el calendario.\n'
             '3. Pulsa «Publicar disponibilidad».',
             'Los tres días quedan publicados. Si abres la ficha pública de María, esos días ya '
             'aparecen disponibles para reservar.'),

            ('CP-38', 'Retirar disponibilidad',
             'Selecciona un día publicado SIN cita y pulsa «Retirar».',
             'Desaparece del calendario público.'),

            ('CP-39', 'No se retira un día con cita',
             'Intenta seleccionar y retirar el día que reservó Juliana.',
             'NEGATIVO. Ese día no se puede seleccionar, o avisa que tiene una cita. Cancelar una '
             'reserva es otra operación, no se hace desde aquí.'),

            ('CP-40', 'Editar el perfil',
             'Entra a «Mi perfil», cambia el titular y guarda.',
             'Aparece «Cambios guardados». El cambio se ve en la ficha pública.'),
        ],
    ),
    (
        '8. Administración',
        'admin@eterclack.test',
        [
            ('CP-41', 'Ver los indicadores',
             'Ingresa como admin.',
             'Cuatro indicadores: clientes activos, fotógrafos publicados, pendientes de revisión '
             '(en amarillo si hay) y correos fallidos.'),

            ('CP-42', 'Revisar un pendiente',
             'En la pestaña «Pendientes», mira la ficha de Carlos Duarte.',
             'Se ve su correo, titular, biografía, especialidades, zonas, cuántas imágenes tiene '
             'y cuándo se postuló.'),

            ('CP-43', 'Aprobar un fotógrafo',
             'Pulsa «Aprobar» en Carlos Duarte.',
             'Sale de pendientes. Si buscas «Carlos» en la búsqueda pública, AHORA sí aparece.'),

            ('CP-44', 'Rechazo sin motivo suficiente',
             'En otro fotógrafo, pulsa «Rechazar» y escribe solo «no».',
             'NEGATIVO. Exige al menos 10 caracteres. El fotógrafo tiene derecho a saber qué corregir.'),

            ('CP-45', 'Rechazar con motivo',
             'Escribe un motivo real y confirma.',
             'Pasa a «Rechazados». En local, le llega un correo con el motivo textual.'),

            ('CP-46', 'El rechazado ve el motivo',
             'Cierra sesión, ingresa como sofia@eterclack.test.',
             'En su panel ve el estado «Requiere ajustes» con el motivo, y un botón «Volver a '
             'postularme».'),

            ('CP-47', 'Repostularse',
             'Pulsa «Volver a postularme».',
             'Su estado vuelve a «En revisión» y reaparece en la lista de pendientes del admin.'),

            ('CP-48', 'Suspender',
             'Como admin, en «Aprobados», suspende a uno.',
             'Deja de aparecer en la búsqueda pública inmediatamente.'),
        ],
    ),
    (
        '9. Gestión de usuarios',
        'admin@eterclack.test · en Administración → Usuarios',
        [
            ('CP-49', 'Ver el listado',
             'Ingresa como admin y entra a «Usuarios».',
             'Cuatro contadores arriba: total, clientes, fotógrafos y administración. Debajo, la '
             'lista con el rol y el estado de cada uno. Tu propia fila lleva la etiqueta «Tú».'),

            ('CP-50', 'Filtrar por rol',
             'En el filtro de rol elige «Fotógrafo».',
             'Quedan solo los fotógrafos. El contador de la derecha cuadra con lo que se ve.'),

            ('CP-51', 'Buscar por correo',
             'Escribe «juliana» en el buscador y pulsa Enter.',
             'Queda solo Juliana Restrepo.'),

            ('CP-52', 'Crear un CLIENTE',
             '1. Pulsa «Crear usuario».\n'
             '2. Deja el tipo en «Cliente».\n'
             '3. Llena nombre, un correo nuevo y una contraseña de 8+ caracteres.\n'
             '4. Deja marcado «Marcar el correo como verificado».\n'
             '5. Pulsa «Crear usuario».',
             'Aparece en la lista como Cliente · Activo. Cierra sesión, ingresa con ese correo y '
             'esa contraseña: entra y puede reservar sin verificar nada.'),

            ('CP-53', 'Crear un FOTÓGRAFO',
             'Repite CP-52 pero elige el tipo «Fotógrafo».',
             'En su fila aparece «Perfil: nombre-apellido · PENDING». Ingresa con esa cuenta y '
             'entra a «Mi perfil»: el panel del fotógrafo carga. NO aparece todavía en la '
             'búsqueda pública, porque está pendiente de aprobación.'),

            ('CP-54', 'Crear un ADMINISTRADOR',
             'Repite CP-52 pero elige el tipo «Administración».',
             'Al elegir ese tipo sale un aviso en amarillo advirtiendo que tendrá control total. '
             'Ingresa con esa cuenta: puede entrar a Administración y ver el listado de usuarios.'),

            ('CP-55', 'Correo repetido',
             'Intenta crear un usuario con juliana@eterclack.test.',
             'NEGATIVO. «Ese correo ya tiene una cuenta.» No se crea nada.'),

            ('CP-56', 'Contraseña corta',
             'Intenta crear un usuario con una contraseña de 5 caracteres.',
             'NEGATIVO. El mensaje aparece bajo el campo de contraseña, no como error general.'),

            ('CP-57', 'Editar un usuario',
             'Pulsa «Editar» en cualquiera, cambia el nombre y guarda.',
             'Aparece «Cambios guardados» y el nombre nuevo se ve en la lista.'),

            ('CP-58', 'Suspender corta la sesión',
             '1. Ingresa como Daniel en otro navegador.\n'
             '2. Como admin, pulsa «Suspender» en Daniel.\n'
             '3. Recarga en el navegador de Daniel.',
             'Daniel queda desconectado de inmediato. Suspender revoca las sesiones abiertas, no '
             'solo impide entrar de nuevo.'),

            ('CP-59', 'Cambiar la contraseña de alguien',
             '1. Pulsa «Contraseña» en un usuario.\n'
             '2. Escribe una nueva y confirma.\n'
             '3. Ingresa con esa cuenta y la contraseña nueva.',
             'Funciona. La contraseña se muestra en claro a propósito, para poder copiarla y '
             'dársela a esa persona.'),

            ('CP-60', 'Borrar y restaurar',
             '1. Borra un usuario sin reservas.\n'
             '2. Marca «Incluir borrados».\n'
             '3. Pulsa «Restaurar».',
             'Al borrarlo desaparece de la lista normal. Con «Incluir borrados» reaparece en gris '
             'con la etiqueta «Borrado». Al restaurarlo vuelve a Activo.'),

            ('CP-61', 'No puedes borrarte a ti mismo',
             'Busca tu propia fila (la que dice «Tú») y mira los botones.',
             'NEGATIVO. «Suspender» y «Borrar» están deshabilitados. Al pasar el cursor explican '
             'por qué. Es lo que evita que te dejes fuera de tu propia plataforma.'),

            ('CP-62', 'No puedes quitarte tu propio rol',
             'Pulsa «Editar» en tu propia fila y mira el campo Rol.',
             'NEGATIVO. El selector está deshabilitado con la nota «No puedes cambiar tu propio rol».'),

            ('CP-63', 'Un fotógrafo con reservas no cambia de rol',
             'Edita a María Gómez (tiene reservas) y cambia su rol a Cliente. Guarda.',
             'NEGATIVO. «Tiene N reserva(s). Suspéndelo en vez de cambiarle el rol.» Cambiarlo '
             'dejaría sus reservas sin fotógrafo.'),

            ('CP-64', 'Correo de un usuario borrado',
             '1. Borra un usuario.\n'
             '2. Intenta crear otro con ese mismo correo.',
             'NEGATIVO. «Ese correo pertenece a un usuario borrado. Restáuralo en vez de crear '
             'otro.» El mensaje dice qué hacer, no solo que no se puede.'),

            ('CP-65', 'Un cliente no entra a Usuarios',
             'Como juliana@eterclack.test, escribe la dirección .../admin/usuarios',
             'NEGATIVO. Te devuelve al inicio. No debe verse ni un instante el listado.'),
        ],
    ),
    (
        '10. Seguridad',
        'Estos casos DEBEN fallar. Si alguno pasa, hay una fuga de datos.',
        [
            ('CP-66', 'Ver la orden de otro cliente',
             '1. Como Juliana, copia la dirección de tu orden (…/ordenes/xxxx).\n'
             '2. Cierra sesión, ingresa como daniel@eterclack.test.\n'
             '3. Pega esa dirección.',
             'NEGATIVO CRÍTICO. No debe mostrar los datos. Debe dar error de permiso.'),

            ('CP-67', 'Cliente entra al panel de admin',
             'Como Juliana, escribe la dirección .../admin',
             'NEGATIVO. Te devuelve al inicio. No debe verse ni un instante el contenido.'),

            ('CP-68', 'Cliente entra al panel de fotógrafo',
             'Como Juliana, escribe .../panel/agenda',
             'NEGATIVO. Te devuelve al inicio.'),

            ('CP-69', 'Fotógrafo entra al panel de admin',
             'Como María, escribe .../admin',
             'NEGATIVO. Te devuelve al inicio.'),

            ('CP-70', 'Fotógrafo acepta el contrato del cliente',
             'Como María, abre la orden de Juliana.',
             'Puede VER la orden (es suya también) pero NO aparece el formulario para aceptar el '
             'contrato. Eso solo lo hace el cliente.'),

            ('CP-71', 'Cambiar contraseña cierra las sesiones',
             '1. Ingresa como Juliana en dos navegadores.\n'
             '2. En uno, cambia la contraseña por recuperación.\n'
             '3. Recarga en el otro.',
             'El segundo navegador queda desconectado. Cambiar la contraseña revoca todas las '
             'sesiones. Solo se puede probar en local: requiere correo.'),
        ],
    ),
    (
        '11. Móvil y app instalable',
        'Con un teléfono real, no el emulador del navegador',
        [
            ('CP-72', 'Navegación en móvil',
             'Abre la aplicación en el teléfono.',
             'El menú aparece plegado (icono de tres líneas). Al abrirlo se ven las opciones de tu rol.'),

            ('CP-73', 'Nada se sale de la pantalla',
             'Recorre todas las pantallas deslizando.',
             'La página NUNCA se desliza hacia los lados. Solo hacia abajo.'),

            ('CP-74', 'Todo se puede tocar',
             'Intenta pulsar los enlaces del pie, los filtros y los días del calendario con el dedo.',
             'Todo se pulsa a la primera. Nada exige precisión de uña.'),

            ('CP-75', 'Reservar desde el móvil',
             'Haz una reserva completa desde el teléfono.',
             'Los tres productos se apilan uno debajo de otro. El calendario se ve completo. '
             'Se puede completar la reserva.'),

            ('CP-76', 'Los campos no hacen zoom',
             'Toca un campo de texto en iPhone.',
             'La pantalla NO hace zoom automático al enfocar.'),

            ('CP-77', 'Instalar la app',
             'En Chrome, menú → «Instalar aplicación». En iPhone, Compartir → «Añadir a inicio».',
             'Queda un icono con la marca. Al abrirlo, la app arranca a pantalla completa, sin '
             'barra de direcciones.'),

            ('CP-78', 'Sin conexión',
             '1. Abre la app instalada.\n'
             '2. Activa el modo avión.\n'
             '3. Recarga.',
             'Carga la estructura de la app, no la pantalla de error del navegador.'),

            ('CP-79', 'Sin conexión NO finge que funciona',
             'En modo avión, intenta hacer una reserva.',
             'NEGATIVO ESPERADO. Debe FALLAR con un error visible. Nunca debe decir que reservó: '
             'una reserva falsa es peor que un error.'),
        ],
    ),
]


def bloque_caso(cid, nombre, pasos, esperado):
    """Un caso: título, pasos/esperado, y el registro para anotar."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f'{cid} · {nombre}')
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DANGER if 'NEGATIVO' in esperado else INK

    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabecera
    fila = t.add_row()
    texto_celda(fila.cells[0], 'PASOS', bold=True, size=8)
    texto_celda(fila.cells[1], 'RESULTADO ESPERADO', bold=True, size=8)
    sombrear(fila.cells[0], LIME_HEX)
    sombrear(fila.cells[1], LIME_HEX)

    # Contenido
    fila = t.add_row()
    texto_celda(fila.cells[0], pasos)
    texto_celda(fila.cells[1], esperado)

    # ── Registro ────────────────────────────────────────────────
    fila = t.add_row()
    celda = fila.cells[0].merge(fila.cells[1])
    celda.text = ''
    p = celda.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    for texto, negrita in [
        ('RESULTADO   ', True),
        (f'{CAJA} OK      {CAJA} Falla      {CAJA} Bloqueado      {CAJA} No aplica', False),
        ('          SEVERIDAD   ', True),
        (f'{CAJA} Bloqueante  {CAJA} Alta  {CAJA} Media  {CAJA} Baja', False),
    ]:
        r = p.add_run(texto)
        r.font.size = Pt(8.5)
        r.bold = negrita
        if negrita:
            r.font.color.rgb = GREY
    sombrear(celda, BAND_HEX)

    # Espacio para escribir
    fila = t.add_row()
    fila.height = Cm(1.5)
    fila.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    celda = fila.cells[0].merge(fila.cells[1])
    celda.text = ''
    p = celda.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run('Observaciones (qué viste exactamente) y qué corregir:')
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = GREY
    sombrear(celda, FIELD_HEX)

    for f in t.rows:
        f.cells[0].width = Cm(8.4)
        if len(f.cells) > 1:
            f.cells[1].width = Cm(9)

    no_partir(t)
    doc.add_paragraph()
    return t


for seccion, quien, casos in CASOS:
    doc.add_page_break()
    titulo(seccion)
    parrafo(quien, gris=True)
    for cid, nombre, pasos, esperado in casos:
        bloque_caso(cid, nombre, pasos, esperado)

# ════════════════════════════════════════════════════════════════
doc.add_page_break()
titulo('11. Resumen de la sesión')

parrafo('Cuenta los casos al terminar. Da el estado real de un vistazo.', gris=True)

t = doc.add_table(rows=0, cols=4)
t.style = 'Table Grid'
fila = t.add_row()
for i, h in enumerate(['OK', 'Falla', 'Bloqueado', 'No aplica']):
    texto_celda(fila.cells[i], h, bold=True)
    sombrear(fila.cells[i], LIME_HEX)
fila = t.add_row()
fila.height = Cm(1.1)
fila.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
for i in range(4):
    texto_celda(fila.cells[i], '')
    fila.cells[i].width = Cm(4.35)
doc.add_paragraph()

parrafo('Veredicto', negrita=True)
p = doc.add_paragraph()
for texto in [
    f'{CAJA} Se puede mostrar al cliente          ',
    f'{CAJA} Solo con las correcciones marcadas          ',
    f'{CAJA} No se puede mostrar todavía',
]:
    r = p.add_run(texto)
    r.font.size = Pt(10)
doc.add_paragraph()

titulo('12. Registro de defectos', 2)
parrafo(
    'Copia aquí SOLO los casos que fallaron. Esta es la lista que se convierte en trabajo de '
    'corrección: si un defecto no está aquí, no se va a arreglar.',
    gris=True,
)

t = doc.add_table(rows=0, cols=5)
t.style = 'Table Grid'
fila = t.add_row()
for i, h in enumerate(['Caso', 'Sev.', 'Qué falló', 'Corrección propuesta', 'Estado']):
    texto_celda(fila.cells[i], h, bold=True, size=8)
    sombrear(fila.cells[i], LIME_HEX)
for _ in range(14):
    fila = t.add_row()
    fila.height = Cm(1.0)
    fila.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for i, w in enumerate([1.7, 1.5, 5.2, 5.6, 3.4]):
        texto_celda(fila.cells[i], '')
        fila.cells[i].width = Cm(w)
doc.add_paragraph()

nota('Estado: pendiente · en curso · corregido · descartado (con motivo).')

# ── Lista de verificación rápida ────────────────────────────────
doc.add_page_break()
titulo('13. Lista de verificación rápida', 2)
parrafo(
    'Para marcar mientras pruebas, sin salir de la pantalla. El detalle va en el bloque de cada caso.',
    gris=True,
)

todos = [(cid, nombre) for _, _, casos in CASOS for cid, nombre, _, _ in casos]
mitad = (len(todos) + 1) // 2

t = doc.add_table(rows=0, cols=4)
t.style = 'Table Grid'
fila = t.add_row()
for i, h in enumerate(['', 'Caso', '', 'Caso']):
    texto_celda(fila.cells[i], h, bold=True, size=8)
    sombrear(fila.cells[i], LIME_HEX)
for i in range(mitad):
    fila = t.add_row()
    izq = todos[i]
    der = todos[i + mitad] if i + mitad < len(todos) else None
    texto_celda(fila.cells[0], CAJA, size=10)
    texto_celda(fila.cells[1], f'{izq[0]} · {izq[1]}', size=8)
    texto_celda(fila.cells[2], CAJA if der else '', size=10)
    texto_celda(fila.cells[3], f'{der[0]} · {der[1]}' if der else '', size=8)
    for j, w in enumerate([0.8, 7.6, 0.8, 7.6]):
        fila.cells[j].width = Cm(w)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
doc.add_page_break()
titulo('14. Qué NO se prueba todavía')
parrafo(
    'Estas funciones no están construidas. Si alguien las busca durante la prueba, no las va a '
    'encontrar, y eso es lo esperado. No las registres como defectos:'
)
tabla(
    ['Función', 'Fase', 'Estado'],
    [
        ['Galerías de fotos y entrega al cliente', '5', 'Sin construir'],
        ['Pago real con Wompi', '6', 'Sin construir — el botón está deshabilitado'],
        ['Dispersión del dinero al fotógrafo', '7', 'Sin construir'],
        ['Servidor de correo propio', '8', 'Sin construir — hoy se usa Mailpit o un relay'],
        ['Despliegue en Hostinger', '9', 'Sin hacer — Render es temporal'],
    ],
    anchos=[8, 2, 7],
)

nota(
    'El estado de la orden llega hasta «Pago pendiente» y ahí se detiene. Es correcto: el cobro '
    'entra en la fase 6.'
)

doc.save('docs/EterClack - Plan de pruebas funcionales.docx')
print('docs/EterClack - Plan de pruebas funcionales.docx')
print(f'  {len(todos)} casos en {len(CASOS)} secciones, cada uno con bloque de registro')

